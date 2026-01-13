"""
Export utilities for converting notes to various formats
"""

from pathlib import Path
from typing import Optional
import markdown
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class PDFExporter:
    """Export content to PDF format"""
    
    def __init__(self):
        self.pdf = FPDF()
    
    def export(self, content: str, output_path: str, title: str = "Lecture Notes"):
        """
        Export content to PDF
        
        Args:
            content: Text content to export
            output_path: Path to save PDF
            title: Document title
        """
        self.pdf.add_page()
        self.pdf.set_font("Arial", "B", 16)
        
        # Add title
        self.pdf.cell(0, 10, title, ln=True, align='C')
        self.pdf.ln(10)
        
        # Add content
        self.pdf.set_font("Arial", "", 12)
        
        # Split content into lines and add to PDF
        for line in content.split('\n'):
            # Handle different markdown elements
            if line.startswith('# '):
                self.pdf.set_font("Arial", "B", 14)
                self.pdf.cell(0, 10, line[2:], ln=True)
                self.pdf.set_font("Arial", "", 12)
            elif line.startswith('## '):
                self.pdf.set_font("Arial", "B", 13)
                self.pdf.cell(0, 10, line[3:], ln=True)
                self.pdf.set_font("Arial", "", 12)
            elif line.startswith('### '):
                self.pdf.set_font("Arial", "B", 12)
                self.pdf.cell(0, 10, line[4:], ln=True)
                self.pdf.set_font("Arial", "", 12)
            else:
                # Regular text
                self.pdf.multi_cell(0, 10, line)
        
        # Save PDF
        self.pdf.output(output_path)


class DOCXExporter:
    """Export content to Word document format"""
    
    def __init__(self):
        self.doc = Document()
    
    def export(self, content: str, output_path: str, title: str = "Lecture Notes"):
        """
        Export content to DOCX
        
        Args:
            content: Text content to export
            output_path: Path to save DOCX
            title: Document title
        """
        # Add title
        title_para = self.doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        for line in content.split('\n'):
            if line.startswith('# '):
                self.doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                self.doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip():
                self.doc.add_paragraph(line)
            else:
                # Empty line
                self.doc.add_paragraph()
        
        # Save document
        self.doc.save(output_path)


class MarkdownExporter:
    """Export content to Markdown format"""
    
    @staticmethod
    def export(content: str, output_path: str):
        """
        Export content to Markdown
        
        Args:
            content: Text content to export
            output_path: Path to save MD file
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)


class HTMLExporter:
    """Export content to HTML format"""
    
    @staticmethod
    def export(content: str, output_path: str, title: str = "Lecture Notes"):
        """
        Export content to HTML
        
        Args:
            content: Markdown content to export
            output_path: Path to save HTML file
            title: Document title
        """
        # Convert markdown to HTML
        html_content = markdown.markdown(content, extensions=['extra', 'codehilite'])
        
        # Create full HTML document
        html_doc = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        h1, h2, h3 {{
            color: #333;
        }}
        h1 {{
            border-bottom: 2px solid #333;
            padding-bottom: 10px;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 5px;
            border-radius: 3px;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            padding-left: 15px;
            color: #666;
        }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_doc)
